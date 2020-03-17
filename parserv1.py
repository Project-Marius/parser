import docx, json, re
from enum import Enum

#Expects a list of two paragraph objects
#Note: assumes that the runs are whitespace-delimited
def parseACFTossup(pgs):
    assert len(pgs) == 2 or len(pgs) == 3
    runs = pgs[0].runs   
    texts = []
    cleanTexts = []
    for run in runs:
        runText = run.text
        cleanTexts.append(runText)
        if run.italic:
            runText = "<i>" + runText + "<\\i>"    
        texts.append(runText)
    allText = "".join(texts)
    cleanText = "".join(cleanTexts)
    number = int(allText.split(".")[0].strip())
    allText = ".".join(allText.split(".")[1:])
    cleanText = ".".join(cleanText.split(".")[1:])
    te = allText.split("(*)")
    assert len(te) == 1 or len(te) == 2
    power = ""
    if len(te) == 2:
        power = te[0].strip()
    rest = te[-1].strip() #Works with or without power
    te = getACFPronounciations(cleanText) 
    clean = te["clean"].strip()
    guides = te["guides"]
    answer = parseACFAnswerLine(pgs[1])
    tags = []
    if len(pgs) == 3:
        tags = parseACFTags(pgs[2])
    return {"number" : number, "clue" : {"power" : power, "non-power" : rest, "clean" : clean}, "guides" : guides, "answer" : answer, "tags" : tags}

def stringToACFGuideRegex(s):
    return re.compile(r"\([\"|" + chr(8220) + r"|" + chr(8221) +  r"]" + s + r"[\"|" + chr(8220) + "|" + chr(8221) + r"]\)")

def getACFPronounciations(s):
    textToUse = s.replace("(*)", "") #In case it's on the border of power
    guideRegex = stringToACFGuideRegex(r"[^\(]*")
    pStrings = guideRegex.findall(textToUse)
    guides = {}
    for guide in pStrings:
        guideText = guide[2:-2]
        numWords = len(guideText.split())
        oldWords = textToUse.split(guide)[0].split()
        assert len(oldWords) >= numWords
        prevWords = oldWords[-numWords:]
        pronouncedText = " ".join(prevWords)
        guides[pronouncedText] = guideText
    return {"guides" : guides, "clean" : guideRegex.sub("", s)}

#For now, just extracts the whole string
#Note: here, we cannot assume that the runs are whitespace delimited
#(i.e. they might split a word)
def parseACFAnswerLine(pg):
    runs = pg.runs
    allTexts = []
    for run in runs:
        runText = run.text
        fs = "" #Totally abusing the definition of a format string
        if run.bold:
            fs += "b"
        if run.italic:
            fs += "i"
        if run.underline:
            fs += "u"
        if len(fs) > 0:
            runText = "<" + fs + ">" + runText + "<\\" + fs + ">"
        allTexts.append(runText)
    tokens = "".join(allTexts).split("[")
    main = tokens[0].strip()
    comments = "[".join(tokens[1:]).strip()
    if len(comments) > 0 and comments[-1] == "]":
        comments = comments[:-1]
    answerRegex = re.compile(r"answer:\s", re.IGNORECASE)
    return {"main" : answerRegex.sub("", main), "comments" : comments}

def parseACFTags(pg):
    text = pg.text.strip()
    if text[0] == "<" and text[-1] == ">":
        text = text[1:-1]
        return [s.strip() for s in text.split(",")]
    else:
        return [s]
    
def parseACFBonusSection(pgs):
    assert len(pgs) == 2
    texts = []
    runs = pgs[0].runs
    for run in runs:
        runText = run.text.strip()
        if run.italic:
            runText = "<i>" + runText + "<\\i>"
        texts.append(runText)
    text = " ".join(texts).strip()
    if text[0:4] == "[10]":
        text = text[4:].strip()
    answer = parseACFAnswerLine(pgs[1])
    return {"clue" : text, "answer" : answer}

def parseACFBonus(pgs):
    assert len(pgs) == 7 or len(pgs) == 8
    introTexts = []
    runs = pgs[0].runs
    for run in runs:
        runText = run.text.strip()
        if run.italic:
            runText = "<i>" + runText + "<\\i>"
        introTexts.append(runText)
    introText = " ".join(introTexts).strip()
    number = int(introText.split(".")[0].strip())
    intro = ".".join(introText.split(".")[1:])
    sections = [parseACFBonusSection(pgs[i:i+2]) for i in range(1, 7, 2)]
    allBonusTexts = []
    for i in [0,1,3,5]:
        allBonusTexts.extend([r.text for r in pgs[i].runs])
    allBonusText = " ".join("".join(allBonusTexts).split()) #Replace all whitespace with 1 space
    guides = getACFPronounciations(allBonusText)["guides"]
    tags = []
    if len(pgs) == 8:
        tags = parseACFTags(pgs[7])
    return {"number" : number, "intro" : intro, "sections" : sections, "guides" : guides, "tags" : tags}

def isEmptyPg(pg):
    return len(pg.text.strip()) == 0

#Adds a run to pg with text text and bold/italic/underline the same as oldRun
def addFormattedRun(pg, text, oldRun):
    newRun = pg.add_run(text)
    newRun.bold = oldRun.bold
    newRun.italic = oldRun.italic
    newRun.underline = oldRun.underline

#Split any paragraphs that contain newlines
def splitParagraphs(pgs):
    newPgs = []
    for pg in pgs:
        if "\n" in pg.text:
            dummyDoc = docx.Document() #A fake document to hold new paragraphs
            newPg = dummyDoc.add_paragraph()
            newPgs.append(newPg)
            for run in pg.runs:
                if "\n" in run.text:
                    tokens = run.text.split("\n")
                    addFormattedRun(newPg, tokens[0], run)
                    if len(tokens) > 2:
                        for token in tokens[1:-1]:
                            newPg = dummyDoc.add_paragraph()
                            newPgs.append(newPg)
                            addFormattedRun(newPg, token, run)
                    newPg = dummyDoc.add_paragraph()
                    newPgs.append(newPg)
                    addFormattedRun(newPg, tokens[-1], run)
                else:
                    addFormattedRun(newPg, run.text, run)
        else:
            newPgs.append(pg)
    return newPgs #Includes empty paragraphs
        
def parseACFFile(name):
    doc = docx.Document(name)
    class State(Enum):
        HEADER = 0
        TOSSUPS = 1
        BONUSES = 2
    state = State.HEADER
    tossups = []
    bonuses = []
    curPgs = []
    splitParas = splitParagraphs(doc.paragraphs)
    for pg in splitParas:
        if "tossups" in pg.text.strip().lower():
            state = State.TOSSUPS
        elif "bonuses" in pg.text.strip().lower():
            state = State.BONUSES
        elif state != State.HEADER:
            if isEmptyPg(pg):
                if len(curPgs) > 0:
                    if state == State.TOSSUPS and len(curPgs) in [2,3]:
                        tossups.append(parseACFTossup(curPgs))
                    elif len(curPgs) in [7,8]:
                        bonuses.append(parseACFBonus(curPgs))
                    else:
                        print("Excluding paragraphs: "[i.text for i in curPgs])
                    curPgs = []
            else:
                curPgs.append(pg)
    jsonString = json.dumps({"tossups" : tossups, "bonuses" : bonuses})
    return jsonString.encode('latin1').decode('unicode_escape') #Hopefully handle wierd characters
